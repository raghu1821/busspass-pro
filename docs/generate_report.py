
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page Setup ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.3)   # 1.3"
section.right_margin  = Cm(2.54)  # 1"
section.top_margin    = Cm(1.91)  # 0.75"
section.bottom_margin = Cm(1.91)  # 0.75"

# ─── Styles helpers ──────────────────────────────────────────────────────────
styles = doc.styles

def set_run_font(run, size, bold=False, name="Times New Roman", color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_paragraph(text, level_size=16, bold=True, center=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, level_size, bold)
    return p

def add_body(text, bold=False, space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, 12, bold)
    return p

def add_section_heading(number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"{number}  {title}")
    set_run_font(run, 14, bold=True)
    return p

def add_bullet(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + indent)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, 12)
    return p

def page_break():
    doc.add_page_break()

# ─── Helper: add header/footer ───────────────────────────────────────────────
def add_header_footer(section, project_title="BUS PASS MANAGEMENT SYSTEM"):
    # Header
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_h = hp.add_run(project_title)
    set_run_font(run_h, 11, bold=False)

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_fl = fp.add_run("Department of CS&E, BIT  2025-26")
    set_run_font(run_fl, 11)

    # Right-aligned page number in footer using tab
    fp.add_run("\t\t\t\t\t")
    run_pg = fp.add_run()
    set_run_font(run_pg, 11)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_pg._r.append(fldChar1)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar2)

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 1 – FRONT SHEET
# ══════════════════════════════════════════════════════════════════════════════
def add_front_sheet():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("BANGALORE INSTITUTE OF TECHNOLOGY")
    set_run_font(run, 16, bold=True)

    p2 = add_body("(Autonomous Institution under VTU)", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body("K.R. Road, V.V.Puram, Bangalore – 560 004", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()  # spacing

    add_body("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    add_body("DATABASE MANAGEMENT SYSTEM MINI PROJECT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body("BCS403", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    add_body('"BUS PASS MANAGEMENT SYSTEM"', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    add_body("Submitted By", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Student table
    table = doc.add_table(rows=5, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ["USN", "NAME", "BATCH", "PHONE NO", "E-MAIL"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 11, bold=True)

    students = [
        ("1BI24CS153", "PRAKHAR KASHYAP",   "C1", "8544402641", "Prakharkashyap0220512@gmail.com"),
        ("1BI24CS158", "PRIYA H",            "C1", "7892885792", "Priyapriyah350@gmail.com"),
        ("1BI24CS160", "RAGHAVENDRA S K",    "C1", "8105750732", "raghavendrakarkanaller@gmail.com"),
        ("1BI24CS171", "SACHIN AMBALI",      "C2", "6363572469", "Sachinambali07@gmail.com"),
    ]
    for r_idx, student in enumerate(students):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(student):
            row.cells[c_idx].text = val
            for run in row.cells[c_idx].paragraphs[0].runs:
                set_run_font(run, 10)

    doc.add_paragraph()
    add_body("Lab In-Charges:  Prof. SUMA L  &  Prof. ASHWINI T N", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body("Academic Year: 2025-2026", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 2 – CERTIFICATE
# ══════════════════════════════════════════════════════════════════════════════
def add_certificate():
    add_heading_paragraph("BANGALORE INSTITUTE OF TECHNOLOGY", 16)
    add_body("(Autonomous Institution under VTU)", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body("K.R. Road, V.V.Puram, Bangalore – 560 004", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_body("Department of Computer Science & Engineering", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_heading_paragraph("Certificate", 16)
    doc.add_paragraph()
    cert_text = (
        'This is to certify that the implementation of DBMS MINI PROJECT entitled '
        '"BUS PASS MANAGEMENT SYSTEM" has been successfully completed by the following '
        'students of IV Semester B.E. for the partial fulfillment of the requirements '
        'for the Bachelor\'s Degree in Computer Science & Engineering of the '
        'Visvesvaraya Technological University during the academic year 2025-2026.'
    )
    add_body(cert_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_paragraph()

    # Student table in certificate
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].text if cell.paragraphs[0].runs else None
    headers_c = ["USN", "Name"]
    for i, h in enumerate(headers_c):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            set_run_font(run, 11, bold=True)
    students_c = [
        ("1BI24CS153", "PRAKHAR KASHYAP"),
        ("1BI24CS158", "PRIYA H"),
        ("1BI24CS160", "RAGHAVENDRA S K"),
        ("1BI24CS171", "SACHIN AMBALI"),
    ]
    for r_idx, s in enumerate(students_c):
        row = table.rows[r_idx+1]
        row.cells[0].text = s[0]
        row.cells[1].text = s[1]

    doc.add_paragraph()
    doc.add_paragraph()
    add_body("Lab In-Charges:", bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_body("Prof. SUMA L                                                       Prof. ASHWINI T N")
    add_body("Designation                                                         Designation")
    add_body("Department of CS&E                                               Department of CS&E")
    add_body("Bangalore Institute of Technology                         Bangalore Institute of Technology")
    add_body("Bangalore                                                              Bangalore")
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 3 – ACKNOWLEDGEMENT
# ══════════════════════════════════════════════════════════════════════════════
def add_acknowledgement():
    add_heading_paragraph("ACKNOWLEDGEMENT", 16)
    doc.add_paragraph()
    ack = (
        "We would like to express our sincere gratitude to all those who have contributed "
        "to the successful completion of this DBMS Mini Project titled "
        '"BUS PASS MANAGEMENT SYSTEM".'
    )
    add_body(ack)
    doc.add_paragraph()
    add_body("We are deeply grateful to the Head of the Department of Computer Science & Engineering, "
             "Bangalore Institute of Technology, for providing the necessary infrastructure and support.")
    doc.add_paragraph()
    add_body("We extend our heartfelt thanks to our Lab In-Charges, Prof. SUMA L and Prof. ASHWINI T N, "
             "for their invaluable guidance, constant encouragement, and constructive suggestions throughout "
             "the course of this project.")
    doc.add_paragraph()
    add_body("We also thank the teaching and non-teaching staff of the Department of CS&E for their "
             "continuous support and motivation.")
    doc.add_paragraph()
    add_body("Finally, we thank our family and friends for their unwavering moral support and encouragement "
             "during the course of this project.")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    add_body("PRAKHAR KASHYAP  (1BI24CS153)", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_body("PRIYA H  (1BI24CS158)", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_body("RAGHAVENDRA S K  (1BI24CS160)", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_body("SACHIN AMBALI  (1BI24CS171)", align=WD_ALIGN_PARAGRAPH.RIGHT)
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 4 – TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
def add_toc():
    add_heading_paragraph("CONTENTS", 16)
    doc.add_paragraph()
    toc_items = [
        ("Chapter 1:", "Introduction",                          "1"),
        ("  1.1",      "Problem Statement",                     "2"),
        ("  1.2",      "Front-end and Back-end Tool Details",   "3"),
        ("Chapter 2:", "Back End Design",                       "5"),
        ("  2.1",      "ER Diagram",                            "6"),
        ("  2.2",      "ER to Relational Mapping",              "7"),
        ("  2.3",      "Normalization up to 3NF",               "9"),
        ("Chapter 3:", "Front End Design",                      "12"),
        ("  3.1",      "Screen Layout Design for WebPages",     "13"),
        ("  3.2",      "Connectivity (Frontend and Backend)",   "15"),
        ("Chapter 4:", "Major Modules with Description",        "17"),
        ("Chapter 5:", "Implementation using Python/MySQL",     "21"),
        ("Chapter 6:", "Snapshots",                             "28"),
        ("Chapter 7:", "Applications",                          "34"),
        ("Chapter 8:", "Conclusion",                            "36"),
    ]
    table = doc.add_table(rows=len(toc_items), cols=3)
    for i, (ch, title, pg) in enumerate(toc_items):
        row = table.rows[i]
        row.cells[0].text = ch
        row.cells[1].text = title
        row.cells[2].text = pg
        for c in range(3):
            for run in row.cells[c].paragraphs[0].runs:
                set_run_font(run, 12, bold=(i in [0,3,6,9,12,13,14,15]))
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 1 – INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
def add_chapter1():
    add_heading_paragraph("CHAPTER-1", 14, center=True)
    add_heading_paragraph("INTRODUCTION", 16, center=True)
    page_break()

    add_section_heading("1.1", "Introduction")
    intro = (
        "A Bus Pass Management System is a database application designed to manage and maintain "
        "all activities related to bus pass issuance, renewal, passenger records, route management, "
        "and pass verification in transport departments and city bus services. The system helps in "
        "storing passenger information, maintaining pass validity details, handling renewal requests, "
        "and managing route and operator records efficiently.\n\n"
        "Traditionally, bus pass operations were maintained manually using registers and paper-based "
        "forms, which often led to errors, data redundancy, and difficulty in searching records during "
        "audits or disputes. The Bus Pass Management System overcomes these problems by providing a "
        "computerized and automated solution for managing bus pass operations.\n\n"
        "The system stores details such as passenger information, pass type and validity, payment records, "
        "route details, operator information, and renewal requests in a centralized database. It allows "
        "administrators to quickly search for passenger records, update pass status after renewal, and "
        "manage pass requests effectively.\n\n"
        "This system plays an important role in day-to-day public transport operations by helping transport "
        "departments and bus operators provide passes quickly to commuters in need. It also improves "
        "coordination between passengers, operators, and the transport authority while reducing manual work "
        "and maintaining data accuracy."
    )
    add_body(intro)
    doc.add_paragraph()
    add_body("The Bus Pass Management System is useful for:", bold=True)
    for item in [
        "Managing passenger records",
        "Maintaining pass validity and status",
        "Processing pass issuance and renewal requests",
        "Generating reports",
        "Ensuring secure data management",
    ]:
        add_bullet(item)

    doc.add_paragraph()
    add_body("Overall, the Bus Pass Management System provides an efficient, reliable, and secure way to manage "
             "public transport pass activities and helps save valuable time for commuters and administrative staff alike.")

    add_section_heading("1.2", "Problem Statement")
    ps = (
        "Urban and semi-urban commuters across India depend heavily on public bus transport for their daily "
        "travel. Managing bus passes for thousands of passengers — including students, employees, senior "
        "citizens, and daily commuters — involves significant administrative challenges:"
    )
    add_body(ps)
    for item in [
        "Manual application processes for issuing new passes or renewals are time-consuming and error-prone",
        "Lack of real-time validation of pass authenticity leading to misuse and fare evasion",
        "Absence of a centralised database to track pass holders, routes, and expiry records",
        "Poor grievance and complaint handling with no digital audit trail",
        "No digital payment gateway integration, forcing passengers to visit counters in person",
        "Difficulty in generating reports for route-wise passenger load and revenue analysis",
    ]:
        add_bullet(item)

    doc.add_paragraph()
    add_body("Existing systems are either paper-based or fragmented across departments. There is a clear need for "
             "a unified, database-driven Bus Pass Management System that:")
    for item in [
        "Allows passengers to register, apply for, and renew passes online",
        "Enables administrators to issue, verify, and revoke passes through a central dashboard",
        "Tracks route assignments, pass validity, and payment history in a structured database",
        "Sends automated renewal reminders and status alerts to registered users",
    ]:
        add_bullet(item)

    add_section_heading("1.3", "Front-End and Back-End Tool Details")
    tool_data = [
        ("Frontend Framework",  "Python Flask (Jinja2)",       "Dynamic HTML rendering and routing"),
        ("Styling",             "HTML5 & CSS3",                 "Responsive page layout and visual design"),
        ("Charting",            "Chart.js / Matplotlib",        "Monthly bar charts and route-wise pie charts"),
        ("Backend Language",    "Python 3.x",                   "Business logic, query execution, session management"),
        ("Database",            "MySQL 8.0",                    "Persistent storage, triggers, stored procedures"),
        ("ORM / Connector",     "mysql-connector-python",       "Flask to MySQL database connectivity"),
        ("Development IDE",     "VS Code / PyCharm",            "Code development and debugging"),
        ("Version Control",     "Git & GitHub",                 "Source code management"),
    ]
    table = doc.add_table(rows=len(tool_data)+1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = ["Category", "Tool / Technology", "Purpose"]
    for i, h in enumerate(hdr):
        c = table.rows[0].cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            set_run_font(run, 12, bold=True)
    for r_idx, (cat, tool, purpose) in enumerate(tool_data):
        row = table.rows[r_idx+1]
        row.cells[0].text = cat
        row.cells[1].text = tool
        row.cells[2].text = purpose
        for c in range(3):
            for run in row.cells[c].paragraphs[0].runs:
                set_run_font(run, 11)
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 2 – BACK END DESIGN
# ══════════════════════════════════════════════════════════════════════════════
def add_chapter2():
    add_heading_paragraph("CHAPTER-2", 14, center=True)
    add_heading_paragraph("BACK END DESIGN", 16, center=True)
    page_break()

    add_section_heading("2.1", "Conceptual Database Design (ER Diagram)")
    add_body(
        "The Entity-Relationship (ER) diagram represents the conceptual structure of the Bus Pass "
        "Management System database. It identifies the main entities, their attributes, and the "
        "relationships between them."
    )
    doc.add_paragraph()
    add_body("Entities and Their Attributes:", bold=True)
    entities = [
        ("Passenger", ["passenger_id (PK)", "full_name", "email (UNIQUE)", "phone", "address", "category", "password_hash", "registered_at"]),
        ("Route",     ["route_id (PK)", "route_name", "source", "destination", "base_fare", "distance_km"]),
        ("Pass_Application", ["application_id (PK)", "passenger_name (FK)", "route_id (FK)", "category", "pass_type", "application_date", "status", "final_fare"]),
        ("Pass",      ["pass_id (PK)", "passenger_name (FK)", "route_id (FK)", "start_date", "end_date", "pass_type", "category", "status", "qr_code_data"]),
        ("Payment",   ["payment_id (PK)", "application_id (FK)", "amount", "payment_mode", "payment_date", "transaction_ref"]),
        ("Concession_Category", ["category_id (PK)", "category_name (UNIQUE)", "concession_rate", "description"]),
        ("Alert",     ["alert_id (PK)", "passenger_name (FK)", "message", "alert_date", "is_read"]),
        ("Admin",      ["admin_id (PK)", "username (UNIQUE)", "password"]),
    ]
    for ename, attrs in entities:
        add_body(f"• {ename}:", bold=True, space_after=2)
        for a in attrs:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"  – {a}")
            set_run_font(run, 11)

    doc.add_paragraph()
    add_body("[ Insert ER Diagram Image Here ]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_body("Key Relationships:", bold=True)
    rels = [
        "A Passenger APPLIES FOR zero or more Pass_Applications (1:N)",
        "A Route is associated with zero or more Pass_Applications (1:N)",
        "A Pass_Application GENERATES zero or one Pass (1:1, trigger-driven)",
        "A Pass_Application has one Payment record (1:1)",
        "A Passenger can receive many Alerts (1:N)",
        "A Passenger can have many Payment records (1:N via Pass_Application)",
        "Concession_Category sets the discount rate for a category of Passenger",
    ]
    for r in rels:
        add_bullet(r)
    page_break()

    add_section_heading("2.2", "Logical Database Design (ER Mapping)")
    add_body(
        "The conceptual ER diagram is mapped to relational tables following standard ER-to-Relational "
        "mapping rules. Each entity becomes a table, and each relationship is represented using foreign keys."
    )
    doc.add_paragraph()
    tables_info = [
        ("Passenger", "passenger_id INT PK AUTO_INCREMENT\nfull_name VARCHAR(100) NOT NULL\nemail VARCHAR(150) UNIQUE NOT NULL\nphone VARCHAR(15)\naddress TEXT\ncategory ENUM('Student','Employee','Senior Citizen','General') NOT NULL\npassword_hash VARCHAR(255) NOT NULL\nregistered_at DATETIME DEFAULT CURRENT_TIMESTAMP"),
        ("Route", "route_id INT PK AUTO_INCREMENT\nroute_name VARCHAR(100) NOT NULL\nsource VARCHAR(100)\ndestination VARCHAR(100)\nbase_fare DECIMAL(10,2)\ndistance_km FLOAT"),
        ("Pass_Application", "application_id INT PK AUTO_INCREMENT\npassenger_name VARCHAR(100) FK → Passenger(full_name)\nroute_id INT FK → Route(route_id)\ncategory VARCHAR(50)\npass_type ENUM('Monthly','Quarterly','Annual')\napplication_date DATE\nstatus ENUM('Pending','Approved','Rejected')\nfinal_fare DECIMAL(10,2)"),
        ("Pass", "pass_id INT PK AUTO_INCREMENT\npassenger_name VARCHAR(100) FK → Passenger(full_name)\nroute_id INT FK → Route(route_id)\nstart_date DATE\nend_date DATE CHECK(end_date > start_date)\npass_type ENUM('Monthly','Quarterly','Annual')\ncategory VARCHAR(50)\nstatus ENUM('Active','Expired','Revoked')\nqr_code_data TEXT"),
        ("Payment", "payment_id INT PK AUTO_INCREMENT\napplication_id INT FK → Pass_Application(application_id)\namount DECIMAL(10,2)\npayment_mode ENUM('Online','Cash','UPI','Card')\npayment_date DATE\ntransaction_ref VARCHAR(100)"),
        ("Concession_Category", "category_id INT PK AUTO_INCREMENT\ncategory_name VARCHAR(50) UNIQUE NOT NULL\nconcession_rate DECIMAL(5,2)\ndescription TEXT"),
        ("Alert", "alert_id INT PK AUTO_INCREMENT\npassenger_name VARCHAR(100) FK → Passenger(full_name)\nmessage TEXT\nalert_date DATE\nis_read BOOLEAN DEFAULT FALSE"),
        ("Admin", "admin_id INT PK AUTO_INCREMENT\nusername VARCHAR(100) UNIQUE NOT NULL\npassword VARCHAR(100) NOT NULL"),
    ]
    for tname, cols in tables_info:
        add_body(f"Table: {tname}", bold=True, space_before=8, space_after=2)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(cols)
        set_run_font(run, 10)
    page_break()

    add_section_heading("2.3", "Normalization up to 3NF with Justification")
    add_body("All tables in the Bus Pass Management System database are normalized up to Third Normal Form (3NF).", bold=False)
    doc.add_paragraph()

    nf_data = [
        ("1NF (First Normal Form)",
         "All attributes contain only atomic (indivisible) values.\n"
         "Each table has a primary key ensuring row uniqueness.\n"
         "Example: The Passenger table stores address as a single TEXT field and phone as VARCHAR(15). "
         "There are no repeating groups or multi-valued attributes.",
         True),
        ("2NF (Second Normal Form)",
         "All non-key attributes are fully functionally dependent on the entire primary key.\n"
         "Since all tables use single-column surrogate primary keys (INT AUTO_INCREMENT), "
         "there are no partial dependencies by definition. Every non-key attribute depends on the full primary key.",
         True),
        ("3NF (Third Normal Form)",
         "There are no transitive dependencies — no non-key attribute depends on another non-key attribute.\n"
         "The Concession_Category table was specifically extracted from Pass_Application to eliminate "
         "the transitive dependency: application_id → category → concession_rate.\n"
         "Route details (source, destination, base_fare) are in the Route table, not duplicated in Pass or Pass_Application.",
         True),
    ]
    for title, desc, _ in nf_data:
        add_body(title, bold=True, space_before=8)
        add_body(desc)

    doc.add_paragraph()
    add_body("Fare Calculation Logic:", bold=True)
    add_body("final_fare = base_fare × (1 − concession_rate)", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body(
        "When a passenger applies for a pass, Flask fetches the base_fare from the Route table based on "
        "the selected route and the concession_rate from the Concession_Category table based on the "
        "passenger category. The result is computed at the application layer and stored in the "
        "Pass_Application table as final_fare.\n\n"
        "Example: Student applies for Monthly pass on Route 500C (base_fare = ₹600), concession 50%  →  "
        "final_fare = 600 × (1 − 0.50) = ₹300"
    )

    add_section_heading("2.4", "Triggers")
    trigger_data = [
        ("trg_renewal_reminder",  "AFTER UPDATE ON Pass",            "Automatically inserts a renewal alert into the Alert table when pass expiry is within 7 days"),
        ("trg_pass_status_update","AFTER UPDATE ON Pass_Application", "Automatically creates a Pass record when application status changes to 'Approved'"),
    ]
    table = doc.add_table(rows=len(trigger_data)+1, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(["Trigger Name", "Event", "Purpose"]):
        c = table.rows[0].cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            set_run_font(run, 12, bold=True)
    for r_idx, row_data in enumerate(trigger_data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx+1].cells[c_idx].text = val
            for run in table.rows[r_idx+1].cells[c_idx].paragraphs[0].runs:
                set_run_font(run, 11)

    add_section_heading("2.5", "Stored Procedures")
    proc_data = [
        ("GetPassengerHistory(passenger_id)",        "Returns complete pass application history, payment records, and active pass details for a given passenger"),
        ("GetRouteSummary(route_id, month, year)",   "Returns total passes issued, revenue collected, and category-wise passenger count for a route in a given month"),
    ]
    table2 = doc.add_table(rows=len(proc_data)+1, cols=2)
    table2.style = 'Table Grid'
    for i, h in enumerate(["Procedure Name", "Purpose"]):
        c = table2.rows[0].cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            set_run_font(run, 12, bold=True)
    for r_idx, (pname, pdesc) in enumerate(proc_data):
        table2.rows[r_idx+1].cells[0].text = pname
        table2.rows[r_idx+1].cells[1].text = pdesc
        for c in range(2):
            for run in table2.rows[r_idx+1].cells[c].paragraphs[0].runs:
                set_run_font(run, 11)
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 3 – FRONT END DESIGN
# ══════════════════════════════════════════════════════════════════════════════
def add_chapter3():
    add_heading_paragraph("CHAPTER-3", 14, center=True)
    add_heading_paragraph("FRONT END DESIGN", 16, center=True)
    page_break()

    add_section_heading("3.1", "Screen Layout Design for WebPages & Forms")
    screens = [
        ("Login / Registration Page",
         "Allows passengers and the admin to log into the system using their email and password. "
         "New passengers can register by providing their full name, email, phone, address, and pass category."),
        ("Passenger Dashboard",
         "After login, passengers can view their current active pass (with QR code), check pass expiry, "
         "submit a new application, track application status, and view renewal alerts."),
        ("Apply for Pass Form",
         "Passengers select a route, pass type (Monthly/Quarterly/Annual), and the system automatically "
         "computes and displays the final fare based on the passenger category concession."),
        ("Admin Dashboard",
         "The admin panel provides statistics on total passengers, active passes, pending applications, "
         "revenue collected, and real-time graphs for monthly pass issuance trends."),
        ("Manage Applications",
         "Admins view all pending applications and can Approve or Reject each with a single click. "
         "Approving triggers the automatic creation of a Pass record (via database trigger)."),
        ("Manage Routes",
         "Admins can add new routes, update route fares, and delete existing routes from the system."),
        ("QR Verification Page",
         "Bus conductors or inspectors can scan/enter a QR code to instantly validate a passenger's pass "
         "and view its status (Active/Expired/Revoked)."),
        ("Payment History Page",
         "Passengers can view all their payment transactions, amounts paid, payment mode, and transaction "
         "references for each pass issued. Admins can view a consolidated payment report."),
    ]
    for title, desc in screens:
        add_body(f"• {title}:", bold=True, space_before=6, space_after=2)
        add_body(f"  {desc}", space_before=0, space_after=4)
    page_break()

    add_section_heading("3.2", "Connectivity (Frontend and Backend)")
    add_body(
        "The frontend (HTML/CSS rendered by Jinja2) communicates with the backend (Python Flask) through "
        "HTTP GET and POST requests. Flask routes handle each action, execute SQL queries using "
        "mysql-connector-python, and return rendered HTML templates with data."
    )
    doc.add_paragraph()
    conn_flow = [
        ("User Action",         "Submits a form or clicks a link"),
        ("HTTP Request",        "Browser sends GET/POST to Flask route in app.py"),
        ("Flask Route Handler", "Validates session, executes SQL via MySQL connector"),
        ("MySQL Database",      "Returns query results / confirms DML operation"),
        ("Jinja2 Template",     "Flask renders HTML template with data variables"),
        ("HTTP Response",       "Browser displays the updated page to the user"),
    ]
    table = doc.add_table(rows=len(conn_flow)+1, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(["Step", "Action"]):
        c = table.rows[0].cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            set_run_font(run, 12, bold=True)
    for r_idx, (step, action) in enumerate(conn_flow):
        table.rows[r_idx+1].cells[0].text = step
        table.rows[r_idx+1].cells[1].text = action
        for c in range(2):
            for run in table.rows[r_idx+1].cells[c].paragraphs[0].runs:
                set_run_font(run, 11)
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 4 – MAJOR MODULES
# ══════════════════════════════════════════════════════════════════════════════
def add_chapter4():
    add_heading_paragraph("CHAPTER-4", 14, center=True)
    add_heading_paragraph("MAJOR MODULES WITH DESCRIPTION", 16, center=True)
    page_break()

    modules = [
        ("4.1", "Passenger Management Module",
         "This module manages all passenger-related operations.",
         ["Add new passenger details during registration",
          "Update passenger profile information (phone, address)",
          "Search passenger records by name, USN, or email",
          "Maintain passenger contact and address details",
          "Secure password hashing using Werkzeug"]),
        ("4.2", "Pass Issuance Module",
         "Handles the lifecycle of bus pass creation.",
         ["Issue new bus passes upon application approval (automated via trigger)",
          "Store pass type (Monthly/Quarterly/Annual) and validity period",
          "Link passenger with pass records using full_name as reference",
          "Generate unique QR code data for each pass using UUID"]),
        ("4.3", "Pass Validity Management Module",
         "Tracks and manages pass status over time.",
         ["Maintain active and expired pass details",
          "Update pass status after renewal",
          "Deactivate pass automatically after expiry",
          "CHECK constraint ensures end_date > start_date"]),
        ("4.4", "Route Management Module",
         "Manages bus routes available in the system.",
         ["Add new routes with source, destination, distance, and base fare",
          "Update base fare for existing routes",
          "Delete routes that are no longer active",
          "Associate passes and applications with specific routes"]),
        ("4.5", "Renewal Request Module",
         "Handles pass renewal workflow.",
         ["Passengers can raise renewal requests from the dashboard",
          "Admin reviews and approves/rejects renewal applications",
          "Track renewal request status (Pending/Approved/Rejected)",
          "System automatically creates a new pass upon approval"]),
        ("4.6", "Payment Module",
         "Records all financial transactions.",
         ["Record payment details for pass issuance and renewal",
          "Store payment mode (Online/Cash/UPI/Card) and date",
          "Link each payment with the corresponding Pass_Application",
          "Store transaction reference numbers for audit"]),
        ("4.7", "Admin Module",
         "Central control panel for administrators.",
         ["Secure admin login with session management",
          "Manage all passengers, passes, routes, and applications",
          "View statistics: total passes issued, revenue, pending requests",
          "Bar charts showing monthly pass issuance trends",
          "Pie charts for route-wise and category-wise distribution"]),
        ("4.8", "Document Verification Module",
         "Handles passenger document upload and admin verification.",
         ["Passengers upload ID proof documents (PDF/image) during registration",
          "Admin views and approves or rejects document submissions",
          "Document status shown on passenger dashboard (Not Uploaded/Pending/Verified)",
          "Only verified passengers can apply for concession-based pass fares"]),
        ("4.9", "QR Code Verification Module",
         "Validates passes at the point of travel.",
         ["Each approved pass carries a unique QR code (UUID-based)",
          "Conductors can enter the QR code to verify pass validity",
          "System displays passenger name, route, validity, and status",
          "Immediately identifies Expired or Revoked passes"]),
    ]
    for num, title, desc, bullets in modules:
        add_section_heading(num, title)
        add_body(desc)
        for b in bullets:
            add_bullet(b)
        doc.add_paragraph()
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 5 – IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
def add_chapter5():
    add_heading_paragraph("CHAPTER-5", 14, center=True)
    add_heading_paragraph("IMPLEMENTATION USING PYTHON / MYSQL", 16, center=True)
    page_break()

    add_section_heading("5.1", "Database Creation (SQL)")
    add_body("The following SQL statements create the database schema for the Bus Pass Management System:")
    doc.add_paragraph()

    sql_blocks = [
        ("Passenger Table",
"""CREATE TABLE Passenger (
    passenger_id INT AUTO_INCREMENT PRIMARY KEY,
    full_name    VARCHAR(100) NOT NULL,
    email        VARCHAR(150) UNIQUE NOT NULL,
    phone        VARCHAR(15),
    address      TEXT,
    category     ENUM('Student','Employee','Senior Citizen','General') NOT NULL,
    password_hash VARCHAR(255) NOT NULL,
    registered_at DATETIME DEFAULT CURRENT_TIMESTAMP
);"""),
        ("Route Table",
"""CREATE TABLE Route (
    route_id     INT AUTO_INCREMENT PRIMARY KEY,
    route_name   VARCHAR(100) NOT NULL,
    source       VARCHAR(100),
    destination  VARCHAR(100),
    base_fare    DECIMAL(10,2),
    distance_km  FLOAT
);"""),
        ("Concession_Category Table",
"""CREATE TABLE Concession_Category (
    category_id    INT AUTO_INCREMENT PRIMARY KEY,
    category_name  VARCHAR(50) UNIQUE NOT NULL,
    concession_rate DECIMAL(5,2),
    description    TEXT
);"""),
        ("Pass_Application Table",
"""CREATE TABLE Pass_Application (
    application_id   INT AUTO_INCREMENT PRIMARY KEY,
    passenger_name   VARCHAR(100),
    route_id         INT,
    category         VARCHAR(50),
    pass_type        ENUM('Monthly','Quarterly','Annual'),
    application_date DATE,
    status           ENUM('Pending','Approved','Rejected') DEFAULT 'Pending',
    final_fare       DECIMAL(10,2),
    FOREIGN KEY (route_id) REFERENCES Route(route_id)
);"""),
        ("Pass Table",
"""CREATE TABLE Pass (
    pass_id        INT AUTO_INCREMENT PRIMARY KEY,
    passenger_name VARCHAR(100),
    route_id       INT,
    start_date     DATE,
    end_date       DATE,
    pass_type      ENUM('Monthly','Quarterly','Annual'),
    category       VARCHAR(50),
    status         ENUM('Active','Expired','Revoked') DEFAULT 'Active',
    qr_code_data   TEXT,
    CONSTRAINT chk_dates CHECK (end_date > start_date),
    FOREIGN KEY (route_id) REFERENCES Route(route_id)
);"""),
        ("Payment Table",
"""CREATE TABLE Payment (
    payment_id     INT AUTO_INCREMENT PRIMARY KEY,
    application_id INT,
    amount         DECIMAL(10,2),
    payment_mode   ENUM('Online','Cash','UPI','Card'),
    payment_date   DATE,
    transaction_ref VARCHAR(100),
    FOREIGN KEY (application_id) REFERENCES Pass_Application(application_id)
);"""),
        ("Alert Table",
"""CREATE TABLE Alert (
    alert_id       INT AUTO_INCREMENT PRIMARY KEY,
    passenger_name VARCHAR(100),
    message        TEXT,
    alert_date     DATE,
    is_read        BOOLEAN DEFAULT FALSE
);"""),
        ("Trigger: trg_renewal_reminder",
"""DELIMITER //
CREATE TRIGGER trg_renewal_reminder
AFTER UPDATE ON Pass
FOR EACH ROW
BEGIN
    IF NEW.end_date <= DATE_ADD(CURDATE(), INTERVAL 7 DAY)
       AND NEW.status = 'Active' THEN
        INSERT INTO Alert (passenger_name, message, alert_date)
        VALUES (NEW.passenger_name,
                CONCAT('Your pass expires on ', NEW.end_date,
                       '. Please renew it.'),
                CURDATE());
    END IF;
END //
DELIMITER ;"""),
        ("Trigger: trg_pass_status_update",
"""DELIMITER //
CREATE TRIGGER trg_pass_status_update
AFTER UPDATE ON Pass_Application
FOR EACH ROW
BEGIN
    IF NEW.status = 'Approved' AND OLD.status != 'Approved' THEN
        INSERT INTO Pass (passenger_name, route_id, start_date,
                          end_date, pass_type, category, qr_code_data)
        VALUES (NEW.passenger_name, NEW.route_id, CURDATE(),
                CASE NEW.pass_type
                    WHEN 'Monthly'   THEN DATE_ADD(CURDATE(), INTERVAL 1 MONTH)
                    WHEN 'Quarterly' THEN DATE_ADD(CURDATE(), INTERVAL 3 MONTH)
                    WHEN 'Annual'    THEN DATE_ADD(CURDATE(), INTERVAL 1 YEAR)
                END,
                NEW.pass_type, NEW.category,
                UUID());
    END IF;
END //
DELIMITER ;"""),
    ]

    for block_title, sql in sql_blocks:
        add_body(block_title, bold=True, space_before=8, space_after=2)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(sql)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    add_section_heading("5.2", "Sample Python Flask Route (app.py)")
    add_body("The following code shows the apply_pass route that handles pass application submission:")
    doc.add_paragraph()
    flask_code = '''@app.route('/apply_pass', methods=['GET', 'POST'])
def apply_pass():
    if 'user' not in session:
        return redirect(url_for('login'))
    conn   = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    if request.method == 'POST':
        route_id   = request.form['route_id']
        pass_type  = request.form['pass_type']
        category   = session['category']
        # Fetch concession rate
        cursor.execute(
            "SELECT concession_rate FROM Concession_Category WHERE category_name=%s",
            (category,))
        conc = cursor.fetchone()
        concession = conc['concession_rate'] if conc else 0
        # Fetch base fare
        cursor.execute(
            "SELECT base_fare FROM Route WHERE route_id=%s", (route_id,))
        route = cursor.fetchone()
        base_fare  = float(route['base_fare']) if route else 0
        final_fare = base_fare * (1 - float(concession))
        # Insert application
        cursor.execute("""
            INSERT INTO Pass_Application
                (passenger_name, route_id, category, pass_type,
                 application_date, status, final_fare)
            VALUES (%s, %s, %s, %s, CURDATE(), 'Pending', %s)
        """, (session['user'], route_id, category, pass_type, final_fare))
        conn.commit()
        flash('Application submitted successfully!', 'success')
        return redirect(url_for('dashboard'))
    # GET: fetch routes
    cursor.execute("SELECT *, base_fare AS fare FROM Route")
    routes = cursor.fetchall()
    return render_template('apply_pass.html', routes=routes)'''

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(flask_code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 6 – SNAPSHOTS
# ══════════════════════════════════════════════════════════════════════════════
def add_chapter6():
    add_heading_paragraph("CHAPTER-6", 14, center=True)
    add_heading_paragraph("SNAPSHOTS", 16, center=True)
    page_break()

    snapshots = [
        "6.1  Login / Registration Page",
        "6.2  Passenger Dashboard with Active Pass and QR Code",
        "6.3  Apply for Pass Form with Auto-Calculated Fare",
        "6.4  Admin Dashboard with Statistics and Charts",
        "6.5  Manage Applications (Approve / Reject)",
        "6.6  Manage Routes Page",
        "6.7  QR Code Verification Page",
        "6.8  Payment History Page",
    ]
    for snap in snapshots:
        add_body(snap, bold=True, space_before=10)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ Insert Screenshot: {snap} ]")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 7 – APPLICATIONS
# ══════════════════════════════════════════════════════════════════════════════
def add_chapter7():
    add_heading_paragraph("CHAPTER-7", 14, center=True)
    add_heading_paragraph("APPLICATIONS", 16, center=True)
    page_break()

    apps = [
        ("Public Transport Pass Management",
         "Transport departments use the system to manage bus pass issuance, renewal, and validation "
         "for daily commuters across various routes and services."),
        ("Passenger Record Management",
         "The system stores passenger details such as name, age, address, contact number, and pass history "
         "for easy access and maintenance."),
        ("Pass Validity Tracking",
         "It helps transport authorities maintain accurate records of active and expired passes and updates "
         "status automatically after renewal or expiry."),
        ("Emergency Pass Allocation",
         "During urgent travel needs or emergencies, the system quickly identifies and issues the required "
         "pass type, saving valuable time for commuters."),
        ("Reduction of Manual Work",
         "The application reduces paperwork and manual calculations by storing all information digitally "
         "in a centralized database."),
        ("Renewal Tracking",
         "It keeps track of previous renewals, renewal dates, and pass validity periods, helping in "
         "maintaining passenger travel history properly."),
        ("Operator and Transport Authority Coordination",
         "The system improves communication between bus operators and transport authorities for faster "
         "pass approvals and renewals."),
        ("Quick Pass Search",
         "Users can easily search for passenger pass details and check validity status instantly through the system."),
        ("Report Generation",
         "The application generates reports related to passengers, pass status, payments, and renewal "
         "requests for administrative purposes."),
        ("Secure Data Management",
         "It provides secure storage of sensitive passenger and payment information and ensures proper "
         "data management and accuracy."),
    ]
    for i, (title, desc) in enumerate(apps, 1):
        add_body(f"{i}. {title}", bold=True, space_before=8, space_after=2)
        add_body(desc, space_before=0, space_after=6)
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 8 – CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
def add_chapter8():
    add_heading_paragraph("CHAPTER-8", 14, center=True)
    add_heading_paragraph("CONCLUSION", 16, center=True)
    page_break()

    add_body(
        "The Bus Pass Management System developed as part of the DBMS Mini Project (BCS403) successfully "
        "demonstrates the power of a well-designed relational database in solving real-world administrative "
        "problems in the public transport domain."
    )
    doc.add_paragraph()
    add_body(
        "The system was designed with a normalized relational database (up to 3NF) using MySQL 8.0, "
        "consisting of 6 interrelated tables: Passenger, Route, Pass_Application, Pass, Payment, "
        "and Admin. The backend was implemented using Python 3.x with the "
        "Flask micro-framework, and the frontend was rendered dynamically using Jinja2 templates with "
        "HTML5 and CSS3."
    )
    doc.add_paragraph()
    add_body("Key achievements of this project include:", bold=True)
    achievements = [
        "Complete end-to-end digital workflow for bus pass application, approval, and issuance",
        "Automated pass creation and renewal alerts using MySQL triggers",
        "QR code-based pass verification for real-time pass validation",
        "Concession-based fare calculation for multiple passenger categories",
        "Comprehensive admin dashboard with statistical graphs and charts",
        "Role-based access control with secure session management",
        "Data integrity enforced through primary keys, foreign keys, NOT NULL, UNIQUE, CHECK constraints",
    ]
    for a in achievements:
        add_bullet(a)

    doc.add_paragraph()
    add_body(
        "The project addresses the limitations of paper-based and fragmented systems by providing a "
        "centralized, database-driven solution that is accurate, efficient, and scalable. Future enhancements "
        "could include integration with real-time payment gateways, mobile applications for commuters, GPS-based "
        "route tracking, and Aadhaar-linked passenger verification."
    )
    doc.add_paragraph()
    add_body(
        "Overall, this project has provided valuable hands-on experience in database design, SQL programming, "
        "web development with Flask, and the application of database management concepts to solve practical "
        "problems in the real world."
    )

# ── Build the Document ──────────────────────────────────────────────────────
add_front_sheet()
add_certificate()
add_acknowledgement()
add_toc()
add_chapter1()
add_chapter2()
add_chapter3()
add_chapter4()
add_chapter5()
add_chapter6()
add_chapter7()
add_chapter8()

# Add header/footer to all sections
for sec in doc.sections:
    add_header_footer(sec)

output_path = r"C:\Users\ragha\OneDrive\Documents\projects\PROJECT\BUSS_PASS_MANAGEMENT\docs\BCS403_Bus_Pass_Management_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
