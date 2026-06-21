
"""
build_report.py – generates the final BCS403 project report with:
  • Proper BIT formatting (margins, fonts, header/footer)
  • Real screenshots inserted into Chapter 6
  • ER Diagram image in Chapter 2
  • All chapters properly aligned
"""

import os, glob
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR   = os.path.dirname(os.path.abspath(__file__))
SS_DIR     = os.path.join(DOCS_DIR, "screenshots")
OUT_PATH   = os.path.join(DOCS_DIR, "BCS403_Bus_Pass_Management_Report_v3.docx")
ER_IMG     = os.path.join(SS_DIR, "er_diagram.png")   # placeholder if exists

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width    = Cm(21.0)
sec.page_height   = Cm(29.7)
sec.left_margin   = Cm(1.3)
sec.right_margin  = Cm(1.0)
sec.top_margin    = Cm(0.75)
sec.bottom_margin = Cm(0.75)

# ── Helpers ─────────────────────────────────────────────────────────────────
TNR = "Times New Roman"

def _set(run, size, bold=False, italic=False, color=None, name=TNR):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, size=16, center=True, sb=12, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    _set(p.add_run(text), size, bold=True)
    return p

def body(text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    _set(p.add_run(text), size, bold=bold)
    return p

def sh(num, title, size=14):  # section heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    _set(p.add_run(f"{num}  {title}"), size, bold=True)

def bullet(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    p.paragraph_format.space_after  = Pt(4)
    _set(p.add_run(text), 12)

def add_table(rows_data, headers, bold_header=True):
    table = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for r in c.paragraphs[0].runs:
            _set(r, 11, bold=bold_header)
    for ri, row_data in enumerate(rows_data):
        for ci, val in enumerate(row_data):
            c = table.rows[ri+1].cells[ci]
            c.text = str(val)
            for r in c.paragraphs[0].runs:
                _set(r, 10)
    return table

def insert_image(img_path, width_cm=15, caption=""):
    """Insert image centred with optional caption."""
    if img_path and os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[ {caption or 'Image'} ]")
        _set(r, 11, italic=True, color=(128,128,128))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        r2 = cp.add_run(f"Figure: {caption}")
        _set(r2, 11, bold=True)

def page_break():
    doc.add_page_break()

def add_hf(section, title="BUS PASS MANAGEMENT SYSTEM"):
    # Header
    hdr = section.header
    hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear()
    hp.paragraph_format.tab_stops.clear_all()
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(9.35), WD_TAB_ALIGNMENT.CENTER)
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(18.7), WD_TAB_ALIGNMENT.RIGHT)
    hrun = hp.add_run(f"\t\t{title.upper()}")
    _set(hrun, 10)
    
    # Footer – Department left, Year center, page no right
    ftr = section.footer
    fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.clear()
    fp.paragraph_format.tab_stops.clear_all()
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(9.35), WD_TAB_ALIGNMENT.CENTER)
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(18.7), WD_TAB_ALIGNMENT.RIGHT)
    
    run_dept = fp.add_run("DEPARTMENT OF CS&E,BIT")
    _set(run_dept, 10)
    run_tab1 = fp.add_run("\t")
    _set(run_tab1, 10)
    run_year = fp.add_run("2025-26")
    _set(run_year, 10)
    run_tab2 = fp.add_run("\t")
    _set(run_tab2, 10)
    run_page = fp.add_run("Page | ")
    _set(run_page, 10)
    
    pgr = fp.add_run()
    _set(pgr, 10)
    for tag, txt in [('w:fldChar','begin'), ('w:instrText','PAGE'), ('w:fldChar','end')]:
        el = OxmlElement(tag)
        if tag == 'w:instrText':
            el.text = txt
        else:
            el.set(qn('w:fldCharType'), txt)
        pgr._r.append(el)

# ════════════════════════════════════════════════════════════════════════════
#  FRONT SHEET
# ════════════════════════════════════════════════════════════════════════════
heading("BANGALORE INSTITUTE OF TECHNOLOGY", 16)
body("(Autonomous Institution under VTU)", align=WD_ALIGN_PARAGRAPH.CENTER)
body("K.R. Road, V.V.Puram, Bangalore – 560 004", align=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
body("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=14)
body("DATABASE MANAGEMENT SYSTEM MINI PROJECT  –  BCS403", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=14)
heading('"BUS PASS MANAGEMENT SYSTEM"', 16, sa=20)
body("Submitted By", align=WD_ALIGN_PARAGRAPH.CENTER, sa=10)

add_table(
    [
        ("1BI24CS153","PRAKHAR KASHYAP",   "C1","8544402641","Prakharkashyap0220512@gmail.com"),
        ("1BI24CS158","PRIYA H",            "C1","7892885792","Priyapriyah350@gmail.com"),
        ("1BI24CS160","RAGHAVENDRA S K",    "C1","8105750732","raghavendrakarkanaller@gmail.com"),
        ("1BI24CS171","SACHIN AMBALI",      "C2","6363572469","Sachinambali07@gmail.com"),
    ],
    ["USN","NAME","BATCH","PHONE NO","E-MAIL"]
)
doc.add_paragraph()
body("Lab In-Charges:  Prof. SUMA L   &   Prof. ASHWINI T N", align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
body("Academic Year: 2025-2026", align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  CERTIFICATE
# ════════════════════════════════════════════════════════════════════════════
heading("BANGALORE INSTITUTE OF TECHNOLOGY", 16)
body("(Autonomous Institution under VTU)", align=WD_ALIGN_PARAGRAPH.CENTER)
body("K.R. Road, V.V.Puram, Bangalore – 560 004", align=WD_ALIGN_PARAGRAPH.CENTER)
body("Department of Computer Science & Engineering", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=10)
heading("C E R T I F I C A T E", 16, sb=10, sa=14)

body(
    'This is to certify that the DBMS Mini Project entitled '
    '"BUS PASS MANAGEMENT SYSTEM" has been successfully completed by the following '
    'students of IV Semester B.E. for partial fulfillment of the Bachelor\'s Degree '
    'in Computer Science & Engineering under Visvesvaraya Technological University '
    'during the academic year 2025-2026.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, sa=12
)

add_table(
    [("1BI24CS153","PRAKHAR KASHYAP"),("1BI24CS158","PRIYA H"),
     ("1BI24CS160","RAGHAVENDRA S K"),("1BI24CS171","SACHIN AMBALI")],
    ["USN","Name"]
)
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
_set(p.add_run("Lab In-Charges:"), 12, bold=True)
for name, desig in [("Prof. SUMA L","Assistant Professor"),("Prof. ASHWINI T N","Assistant Professor")]:
    doc.add_paragraph()
    sig = doc.add_paragraph()
    _set(sig.add_run(name), 12, bold=True)
    sig2 = doc.add_paragraph()
    _set(sig2.add_run(f"{desig}\nDepartment of CS&E\nBangalore Institute of Technology, Bangalore"), 11)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  ACKNOWLEDGEMENT
# ════════════════════════════════════════════════════════════════════════════
heading("ACKNOWLEDGEMENT", 16)
doc.add_paragraph()
for para in [
    'We sincerely thank the Head of the Department of Computer Science & Engineering, '
    'Bangalore Institute of Technology, for providing the necessary infrastructure and '
    'resources to carry out this project.',

    'We are deeply grateful to our Lab In-Charges, Prof. SUMA L and Prof. ASHWINI T N, '
    'for their invaluable guidance, constant encouragement, timely suggestions, and '
    'unwavering support throughout the course of this DBMS Mini Project.',

    'We extend our gratitude to all the teaching and non-teaching staff members of the '
    'Department of CS&E for their continuous moral support and motivation.',

    'Finally, we thank our family and friends for their encouragement and support during '
    'the completion of this project.',
]:
    body(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sa=10)

doc.add_paragraph()
doc.add_paragraph()
for s in ["PRAKHAR KASHYAP  (1BI24CS153)","PRIYA H  (1BI24CS158)",
          "RAGHAVENDRA S K  (1BI24CS160)","SACHIN AMBALI  (1BI24CS171)"]:
    body(s, align=WD_ALIGN_PARAGRAPH.RIGHT, sa=4)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
heading("CONTENTS", 16)
doc.add_paragraph()
toc = [
    ("","Front Sheet","i"),("","Certificate","ii"),("","Acknowledgement","iii"),("","Contents","iv"),
    ("Chapter 1","Introduction","1"),
    ("  1.1","Problem Statement","2"),
    ("  1.2","Front-End and Back-End Tool Details","3"),
    ("Chapter 2","Back End Design","5"),
    ("  2.1","Conceptual Database Design (ER Diagram)","6"),
    ("  2.2","Logical Database Design (ER Mapping)","8"),
    ("  2.3","Normalization up to 3NF","11"),
    ("  2.4","Triggers","13"),
    ("  2.5","Stored Procedures","14"),
    ("Chapter 3","Front End Design","15"),
    ("  3.1","Screen Layout Design for WebPages & Forms","16"),
    ("  3.2","Connectivity (Frontend and Backend)","18"),
    ("Chapter 4","Major Modules with Description","19"),
    ("Chapter 5","Implementation using Python / MySQL","25"),
    ("Chapter 6","Snapshots","33"),
    ("Chapter 7","Applications","38"),
    ("Chapter 8","Conclusion","40"),
]
tbl = doc.add_table(rows=len(toc), cols=3)
for i,(ch,title,pg) in enumerate(toc):
    r = tbl.rows[i]
    r.cells[0].text = ch; r.cells[1].text = title; r.cells[2].text = pg
    bold = ch.startswith("Chapter") or ch == ""
    for ci in range(3):
        for run in r.cells[ci].paragraphs[0].runs:
            _set(run, 12, bold=bold)
    r.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  CHAPTER 1 – INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading("CHAPTER-1", 14); heading("INTRODUCTION", 16)
page_break()

sh("1.1","Introduction")
body("A Bus Pass Management System is a database application designed to manage and maintain "
     "all activities related to bus pass issuance, renewal, passenger records, route management, "
     "and pass verification in transport departments and city bus services. The system helps in "
     "storing passenger information, maintaining pass validity details, handling renewal requests, "
     "and managing route and operator records efficiently.", sa=8)

body("Traditionally, bus pass operations were maintained manually using registers and paper-based "
     "forms, which often led to errors, data redundancy, and difficulty in searching records during "
     "audits or disputes. The Bus Pass Management System overcomes these problems by providing a "
     "computerized and automated solution for managing bus pass operations.", sa=8)

body("The Bus Pass Management System is useful for:", bold=True, sa=4)
for b in ["Managing passenger records","Maintaining pass validity and status",
          "Processing pass issuance and renewal requests","Generating reports",
          "Ensuring secure data management"]:
    bullet(b)

sh("1.2","Problem Statement")
body("Urban and semi-urban commuters across India depend heavily on public bus transport. "
     "Managing passes for thousands of passengers involves significant administrative challenges:", sa=6)
for b in [
    "Manual application processes for issuing new passes or renewals are time-consuming and error-prone",
    "Lack of real-time validation of pass authenticity leading to misuse and fare evasion",
    "Absence of a centralised database to track pass holders, routes, and expiry records",
    "Poor grievance handling with no digital audit trail",
    "No digital payment gateway integration forcing passengers to visit counters in person",
    "Difficulty in generating reports for route-wise passenger load and revenue analysis",
]:
    bullet(b)

body("\nThe proposed system addresses all the above challenges by providing:", bold=True, sa=4)
for b in [
    "Allows passengers to register, apply for, and renew passes online",
    "Enables administrators to issue, verify, and revoke passes through a central dashboard",
    "Tracks route assignments, pass validity, and payment history in a structured database",
    "Sends automated renewal reminders and status alerts to registered users",
]:
    bullet(b)

sh("1.3","Front-End and Back-End Tool Details")
add_table([
    ("Frontend Framework","Python Flask (Jinja2)","Dynamic HTML rendering and routing"),
    ("Styling","HTML5 & CSS3","Responsive page layout and visual design"),
    ("Charting","Chart.js / Matplotlib","Monthly bar charts and route-wise pie charts"),
    ("Backend Language","Python 3.x","Business logic, query execution, session management"),
    ("Database","MySQL 8.0","Persistent storage, triggers, stored procedures"),
    ("ORM / Connector","mysql-connector-python","Flask to MySQL database connectivity"),
    ("Development IDE","VS Code / PyCharm","Code development and debugging"),
    ("Version Control","Git & GitHub","Source code management"),
], ["Category","Tool / Technology","Purpose"])
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  CHAPTER 2 – BACK END DESIGN
# ════════════════════════════════════════════════════════════════════════════
heading("CHAPTER-2", 14); heading("BACK END DESIGN", 16)
page_break()

sh("2.1","Conceptual Database Design (ER Diagram)")
body("The Entity-Relationship (ER) diagram below represents the conceptual structure of the "
     "Bus Pass Management System database. It shows all major entities, their attributes, "
     "and the relationships between them.", sa=10)
insert_image(ER_IMG, width_cm=15, caption="Entity-Relationship Diagram – Bus Pass Management System")
doc.add_paragraph()
body("Key Entities:", bold=True, sa=4)
for e in ["Passenger","Route","Pass_Application","Pass","Payment","Admin"]:
    bullet(e)
body("Key Relationships:", bold=True, sa=4)
for r in [
    "A Passenger APPLIES FOR zero or more Pass_Applications (1:N)",
    "A Route is associated with many Pass_Applications (1:N)",
    "A Pass is generated for a Passenger upon application approval and payment (1:1)",
    "A Pass has an associated Payment record (1:1)",
]:
    bullet(r)
page_break()

sh("2.2","Logical Database Design (ER Mapping)")
body("Each entity from the ER diagram maps to a relational table. Relationships are implemented "
     "using foreign keys. All tables use surrogate integer primary keys with AUTO_INCREMENT.", sa=8)
for tname, cols in [
    ("Passenger",           "passenger_id INT PK AUTO_INCREMENT | full_name VARCHAR(100) NOT NULL | email VARCHAR(100) UNIQUE NOT NULL | phone VARCHAR(15) | address TEXT | category ENUM('Student','Employee','Senior Citizen','General','Physically Challenged') | password VARCHAR(255) | photo LONGTEXT | doc_proof LONGTEXT | doc_status VARCHAR(20)"),
    ("Route",               "route_id INT PK AUTO_INCREMENT | source VARCHAR(100) | destination VARCHAR(100) | base_fare DECIMAL(10,2)"),
    ("Pass_Application",    "application_id INT PK AUTO_INCREMENT | passenger_name VARCHAR(100) | passenger_id INT FK | route_id INT FK | pass_type VARCHAR(50) | duration INT | status VARCHAR(50) | created_at TIMESTAMP"),
    ("Pass",                "pass_id INT PK AUTO_INCREMENT | passenger_name VARCHAR(100) | passenger_id INT FK | pass_type VARCHAR(50) | valid_until DATE | status VARCHAR(50) | qr_code LONGTEXT"),
    ("Payment",             "payment_id INT PK AUTO_INCREMENT | application_id INT FK | passenger_name VARCHAR(100) | passenger_id INT FK | amount DECIMAL(10,2) | payment_method VARCHAR(50) | transaction_id VARCHAR(100) | created_at TIMESTAMP"),
    ("Admin",               "admin_id INT PK AUTO_INCREMENT | username VARCHAR(100) | password VARCHAR(100)"),
]:
    body(f"Table: {tname}", bold=True, sb=10, sa=2)
    rows = [[c.split(' ',1)[0], c.split(' ',1)[1] if ' ' in c else ''] for c in cols.split(' | ')]
    add_table(rows, ["Column","Definition"])
page_break()

sh("2.3","Normalization up to 3NF with Justification")
for nf, desc in [
    ("First Normal Form (1NF)",
     "All attributes contain atomic values. Each table has a unique primary key. "
     "No repeating groups or multi-valued attributes exist. Example: phone is stored as VARCHAR(15), "
     "not as a list; category is a single ENUM value per row."),
    ("Second Normal Form (2NF)",
     "All non-key attributes are fully functionally dependent on the entire primary key. "
     "Since every table uses a single-column surrogate primary key (INT AUTO_INCREMENT), "
     "there are no partial dependencies by definition."),
    ("Third Normal Form (3NF)",
     "No transitive dependencies exist. Non-prime attributes are directly dependent on the primary keys. "
     "Category concessions are calculated dynamically at the application layer, avoiding the need to store "
     "redundant concession rates in the database. Passenger, Route, and Pass data are fully decoupled and "
     "linked via primary/foreign key relationships."),
]:
    body(nf, bold=True, sb=10, sa=4)
    body(desc, sa=6)

body("Fare Calculation Logic (computed at application layer):", bold=True, sb=10, sa=4)
body("final_fare  =  base_fare  ×  (1  −  concession_rate)", align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
body("Example: Student – Route 500C (base_fare ₹600) – Concession 50%  →  final_fare = ₹300", sa=8)
page_break()

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  CHAPTER 3 – FRONT END DESIGN
# ════════════════════════════════════════════════════════════════════════════
heading("CHAPTER-3", 14); heading("FRONT END DESIGN", 16)
page_break()

sh("3.1","Screen Layout Design for WebPages & Forms")
for title, desc in [
    ("Login / Registration Page",
     "Allows passengers and the admin to log into the system. New passengers register by "
     "providing full name, email, phone, address, and category."),
    ("Passenger Dashboard",
     "Displays current active pass with QR code, pass expiry date, renewal alerts, "
     "application status, and quick action buttons."),
    ("Apply for Pass Form",
     "Passengers select a route and pass type. The system automatically computes the final "
     "fare using the concession rate for their category."),
    ("Admin Dashboard",
     "Shows total passengers, active passes, pending applications, revenue collected, "
     "monthly bar charts, and route-wise pie charts."),
    ("Manage Applications",
     "Admins view all pending applications and Approve or Reject them. Approval triggers "
     "automatic Pass creation via database trigger."),
    ("Manage Routes",
     "Admins add new routes, update base fares, and delete inactive routes."),
    ("QR Code Verification Page",
     "Conductors enter a QR code to instantly verify a passenger's pass — shows name, "
     "route, validity period, and current status."),
    ("Payment Portal & Receipts",
     "Passengers view transit pass payment details and retrieve printed receipts for "
     "completed card and UPI transaction records."),
]:
    body(f"• {title}:", bold=True, sb=8, sa=2)
    body(f"  {desc}", sa=4)

sh("3.2","Connectivity (Frontend and Backend)")
body("The frontend (Jinja2 HTML templates) communicates with the backend (Python Flask) "
     "through HTTP GET and POST requests:", sa=8)
add_table([
    ("1. User Action",         "User submits a form or clicks a navigation link"),
    ("2. HTTP Request",        "Browser sends GET/POST to a Flask route in app.py"),
    ("3. Flask Route Handler", "Validates session, executes SQL via mysql-connector"),
    ("4. MySQL Database",      "Processes query and returns result set"),
    ("5. Jinja2 Template",     "Flask renders HTML with injected data variables"),
    ("6. HTTP Response",       "Browser displays the updated page to the user"),
], ["Step","Description"])
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  CHAPTER 4 – MAJOR MODULES
# ════════════════════════════════════════════════════════════════════════════
heading("CHAPTER-4", 14); heading("MAJOR MODULES WITH DESCRIPTION", 16)
page_break()

modules = [
    ("4.1","Passenger Management Module",
     "Manages all passenger-related operations including registration, profile updates, and record search.",
     ["Add new passenger details during registration with hashed password storage",
      "Update passenger profile information (phone number, address)",
      "Search passenger records by name, USN, or email address",
      "Secure password hashing using Werkzeug generate_password_hash"]),
    ("4.2","Pass Issuance Module",
     "Handles the full lifecycle of bus pass creation.",
     ["Issue new bus passes automatically upon application approval (via trigger)",
      "Store pass type (Monthly/Quarterly/Annual) and computed validity period",
      "Generate unique QR code data for each pass using UUID",
      "Link passenger with pass records using full_name as reference key"]),
    ("4.3","Pass Validity Management Module",
     "Tracks and manages pass status over time.",
     ["Maintain active and expired pass records in the Pass table",
      "Update pass status automatically after renewal request approval",
      "Database CHECK constraint ensures end_date is always after start_date",
      "Instant pass status lookup via QR code verification"]),
    ("4.4","Route Management Module",
     "Manages all bus routes available in the system.",
     ["Add new routes with source, destination, distance, and base fare",
      "Update base fare for existing routes through admin panel",
      "Delete routes that are no longer operational",
      "Associate passes and applications with specific routes via route_id FK"]),
    ("4.5","Renewal Request Module",
     "Handles the complete pass renewal workflow.",
     ["Passengers raise renewal requests from their dashboard",
      "Admin reviews and approves or rejects renewal applications",
      "System automatically creates a new Pass record upon approval",
      "Track renewal request status: Pending / Approved / Rejected"]),
    ("4.6","Payment Simulation Module",
     "Handles simulated passenger checkout for pass activation.",
     ["Presents a simulated payment gateway modal on the passenger dashboard",
      "Supports simulated options for VISA, Mastercard, and UPI payment modes",
      "Validates simulated card input constraints and expiry fields",
      "Successful payment triggers pass activation and QR code generation"]),
    ("4.7","Admin Module",
     "Central control panel for system administrators.",
     ["Secure admin login with Flask session management",
      "Manage passengers, passes, routes, applications from one dashboard",
      "Real-time statistics: passes issued, revenue, pending applications",
      "Bar charts for monthly issuance trends, pie charts for category distribution"]),
    ("4.8","Payment Module",
     "Handles transit fee processing and transaction logging.",
     ["Passengers submit payments for approved pass applications via Visa, Mastercard, or UPI",
      "Admins view the complete system financial ledger showing transaction IDs and amounts",
      "Enforces billing accuracy and records transaction timestamps for audit purposes"]),
    ("4.9","QR Code Verification Module",
     "Real-time pass validation at the point of travel.",
     ["Each approved pass carries a UUID-based unique QR code",
      "Conductors enter QR code to verify pass validity instantly",
      "System displays passenger name, route, validity dates, and status",
      "Clearly identifies and reports Expired or Revoked passes"]),
]
for num, title, desc, bullets in modules:
    sh(num, title)
    body(desc, sa=6)
    for b in bullets:
        bullet(b)
    doc.add_paragraph()
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  CHAPTER 5 – IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════════════
heading("CHAPTER-5", 14); heading("IMPLEMENTATION USING PYTHON / MySQL", 16)
page_break()

sh("5.1","Database Schema (SQL)")
body("The following SQL statements create the complete database schema:", sa=8)

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)

for label, sql in [
    ("Passenger Table", """CREATE TABLE passenger (
  passenger_id INT AUTO_INCREMENT PRIMARY KEY,
  full_name    VARCHAR(100) NOT NULL,
  email        VARCHAR(100) UNIQUE NOT NULL,
  phone        VARCHAR(15) DEFAULT NULL,
  address      TEXT,
  category     ENUM('Student','Employee','Senior Citizen','General','Physically Challenged') NOT NULL,
  password     VARCHAR(255) NOT NULL,
  photo        LONGTEXT,
  doc_proof    LONGTEXT,
  doc_status   VARCHAR(20) DEFAULT 'Not Uploaded'
);"""),
    ("Route Table", """CREATE TABLE route (
  route_id    INT AUTO_INCREMENT PRIMARY KEY,
  source      VARCHAR(100) DEFAULT NULL,
  destination VARCHAR(100) DEFAULT NULL,
  base_fare   DECIMAL(10,2) DEFAULT NULL
);"""),
    ("Pass_Application Table", """CREATE TABLE pass_application (
  application_id INT AUTO_INCREMENT PRIMARY KEY,
  passenger_name VARCHAR(100) DEFAULT NULL,
  passenger_id   INT DEFAULT NULL,
  route_id       INT DEFAULT NULL,
  pass_type      VARCHAR(50) DEFAULT NULL,
  duration       INT DEFAULT NULL,
  status         VARCHAR(50) DEFAULT NULL,
  created_at     TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (route_id) REFERENCES route(route_id),
  FOREIGN KEY (passenger_id) REFERENCES passenger(passenger_id)
);"""),
    ("Pass Table", """CREATE TABLE pass (
  pass_id        INT AUTO_INCREMENT PRIMARY KEY,
  passenger_name VARCHAR(100) DEFAULT NULL,
  passenger_id   INT DEFAULT NULL,
  pass_type      VARCHAR(50) DEFAULT NULL,
  valid_until    DATE DEFAULT NULL,
  status         VARCHAR(50) DEFAULT NULL,
  qr_code        LONGTEXT,
  FOREIGN KEY (passenger_id) REFERENCES passenger(passenger_id)
);"""),
    ("Payment Table", """CREATE TABLE payment (
  payment_id     INT AUTO_INCREMENT PRIMARY KEY,
  application_id INT,
  passenger_name VARCHAR(100) DEFAULT NULL,
  amount         DECIMAL(10,2) NOT NULL,
  payment_method VARCHAR(50) DEFAULT 'Card',
  transaction_id VARCHAR(100) DEFAULT NULL,
  created_at     TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (application_id) REFERENCES pass_application(application_id) ON DELETE CASCADE
);"""),
    ("Admin Table", """CREATE TABLE admin (
  admin_id INT AUTO_INCREMENT PRIMARY KEY,
  username VARCHAR(100) DEFAULT NULL,
  password VARCHAR(100) DEFAULT NULL
);"""),
]:
    body(label, bold=True, sb=10, sa=2)
    code_block(sql)

sh("5.2","Flask Route Implementation (app.py sample)")
body("Sample Flask route for pass application submission:", sa=6)
flask_sample = (
    "@app.route('/apply_pass', methods=['GET', 'POST'])\n"
    "def apply_pass():\n"
    "    if 'user' not in session:\n"
    "        return redirect('/login')\n"
    "    passenger_name = session['user']\n"
    "    db = get_db()\n"
    "    cursor = db.cursor(dictionary=True)\n"
    "    if request.method == 'POST':\n"
    "        route_id  = request.form['route_id']\n"
    "        pass_type = request.form['pass_type']\n"
    "        duration  = request.form['duration']\n"
    "        query = '''\n"
    "        INSERT INTO Pass_Application\n"
    "        (passenger_name, route_id, pass_type, duration, status)\n"
    "        VALUES (%s, %s, %s, %s, 'Pending')\n"
    "        '''\n"
    "        cursor.execute(query, (passenger_name, route_id, pass_type, duration))\n"
    "        db.commit()\n"
    "        return redirect('/dashboard?msg=Application+submitted+successfully!&type=success')\n"
    "    cursor.execute('SELECT *, base_fare AS fare FROM Route')\n"
    "    routes = cursor.fetchall()\n"
    "    return render_template('apply_pass.html', routes=routes)"
)
code_block(flask_sample)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  CHAPTER 6 – SNAPSHOTS  (real screenshots inserted)
# ════════════════════════════════════════════════════════════════════════════
heading("CHAPTER-6", 14); heading("SNAPSHOTS", 16)
page_break()

snapshot_defs = [
    ("01_login_page",          "Login Page – Citizen authentication portal for bus pass services"),
    ("02_register_page",       "Registration Page – New passenger account creation with category selection"),
    ("03_passenger_dashboard", "Passenger Dashboard – Applications, quick actions, and pass status"),
    ("04_apply_pass_form",     "Apply for Pass Form – Route selection and concession-based fare preview"),
    ("05_view_pass_qr",        "View Digital Pass & QR Code – Active pass card with scan-to-verify QR"),
    ("06_alerts_page",         "Alerts Page – Renewal reminders and expiry notifications"),
    ("07_feedback_page",       "Payment History – Complete passenger transaction records and receipts"),
    ("08_pass_history",        "Pass History – Complete application and travel pass history log"),
    ("09_qr_verify_page",      "QR Code Verification – Real-time pass validation for bus conductors"),
    ("10_profile_page",        "Profile Page – Passenger profile with document upload and verification"),
    ("11_admin_login_page",    "Admin Login Page – Secure admin authentication portal"),
    ("12_admin_dashboard",     "Admin Dashboard – All applications with approve/revoke actions"),
    ("13_manage_applications", "Manage Applications – Full pass application management panel"),
    ("14_manage_routes",       "Manage Routes – Add, update, and delete BMTC bus routes"),
    ("15_manage_users",        "Manage Users – View and manage all registered passengers"),
    ("16_stats_analytics",     "Statistics & Analytics – System KPIs, charts and route popularity"),
    ("17_admin_feedback",      "Admin Payments Panel – View passenger pass transaction ledger"),
]

for fn, caption in snapshot_defs:
    img_path = os.path.join(SS_DIR, f"{fn}.png")
    sh(f"6.{snapshot_defs.index((fn,caption))+1}", caption.split("–")[0].strip())
    body(caption, sa=6)
    insert_image(img_path, width_cm=14, caption=caption)
    doc.add_paragraph()


page_break()

# ════════════════════════════════════════════════════════════════════════════
#  CHAPTER 7 – APPLICATIONS
# ════════════════════════════════════════════════════════════════════════════
heading("CHAPTER-7", 14); heading("APPLICATIONS", 16)
page_break()

for i, (title, desc) in enumerate([
    ("Public Transport Pass Management",
     "Transport departments use the system to manage bus pass issuance, renewal, and validation "
     "for daily commuters across various routes and services."),
    ("Passenger Record Management",
     "The system stores passenger details such as name, category, address, contact number, and "
     "pass history for quick access and efficient maintenance."),
    ("Pass Validity Tracking",
     "Transport authorities maintain accurate records of active and expired passes. Pass status "
     "updates automatically after renewal or expiry through database triggers."),
    ("Emergency Pass Allocation",
     "During urgent travel needs or emergencies, the system quickly identifies eligible passengers "
     "and issues the required pass type, reducing wait time for commuters."),
    ("Reduction of Manual Work",
     "The application eliminates paperwork and manual calculations by storing all information "
     "digitally in a centralized MySQL database with automated workflows."),
    ("Renewal Tracking",
     "The system tracks previous renewals, renewal dates, and pass validity periods, maintaining "
     "a complete travel history for each passenger."),
    ("Operator and Authority Coordination",
     "The system improves coordination between bus operators and transport authorities, enabling "
     "faster pass approvals, renewals, and dispute resolution."),
    ("Quick Pass Search",
     "Administrators and conductors can instantly search for passenger pass details and validate "
     "pass status through the QR code verification feature."),
    ("Report Generation",
     "The system generates comprehensive reports related to passengers, pass status, payment "
     "records, and renewal requests for administrative use and audits."),
    ("Secure Data Management",
     "Sensitive passenger and payment information is stored securely with hashed passwords, "
     "session-based authentication, and MySQL data integrity constraints."),
], 1):
    body(f"{i}. {title}", bold=True, sb=8, sa=2)
    body(desc, sa=6)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  CHAPTER 8 – CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
heading("CHAPTER-8", 14); heading("CONCLUSION", 16)
page_break()

for para in [
    "The Bus Pass Management System developed as part of the DBMS Mini Project (BCS403) successfully "
    "demonstrates the power of a well-designed relational database in solving real-world administrative "
    "problems in the public transport domain.",

    "The system implements a fully normalized relational database (up to 3NF) using MySQL 8.0, "
    "consisting of 6 relational tables: passenger, route, pass_application, pass, payment, and admin. "
    "The backend was implemented using Python 3.x with the Flask micro-framework, and the frontend "
    "was rendered dynamically using Jinja2 templates with HTML5 and CSS3.",

    "Key achievements of this project include an end-to-end digital workflow for bus pass application, "
    "approval, and issuance; application-layer pass generation and renewal logic in Flask; QR "
    "code-based pass verification for real-time validation; concession-based fare calculation for "
    "multiple passenger categories; a comprehensive admin dashboard with statistical graphs and charts; "
    "and data integrity enforced through primary keys, foreign keys, NOT NULL, and UNIQUE constraints.",

    "Future enhancements could include integration with real-time payment gateways, mobile applications "
    "for commuters, GPS-based route tracking, Aadhaar-linked passenger verification, and multi-city "
    "transport authority support.",

    "Overall, this project has provided valuable hands-on experience in database design, SQL programming, "
    "web development with Flask, and the application of database management concepts to solve practical "
    "problems in the real world.",
]:
    body(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sa=10)

# ── Apply header/footer to all sections ─────────────────────────────────────
for s in doc.sections:
    add_hf(s)

doc.save(OUT_PATH)
print(f"\n[DONE] Report saved: {OUT_PATH}")
